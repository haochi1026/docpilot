#!/usr/bin/env python3
"""Generate deterministic mixed-format documents for the DocPilot benchmark.

Requires Pillow, reportlab and python-docx. The files intentionally contain
different facts so the benchmark verifies retrieval correctness instead of only
checking that an SSE connection stayed open.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas


DOCUMENTS = [
    {
        "filename": "structured-policy.md",
        "kind": "markdown",
        "question": "What is the structured policy release code?",
        "expected": "MD-4821",
    },
    {
        "filename": "operations-table.docx",
        "kind": "docx-table",
        "question": "What is the database recovery exercise code?",
        "expected": "DOCX-5932",
    },
    {
        "filename": "native-policy.pdf",
        "kind": "native-pdf",
        "question": "What is the native PDF audit code?",
        "expected": "PDF-6843",
    },
    {
        "filename": "scanned-policy.pdf",
        "kind": "scanned-pdf-ocr",
        "question": "What is the scanned document emergency code?",
        "expected": "OCR-7319",
    },
]


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--output",
        type=Path,
        default=Path(__file__).resolve().parent / "fixtures" / "performance",
    )
    args = parser.parse_args()
    output = args.output.resolve()
    output.mkdir(parents=True, exist_ok=True)
    make_markdown(output / "structured-policy.md")
    make_docx(output / "operations-table.docx")
    make_native_pdf(output / "native-policy.pdf")
    make_scanned_pdf(output / "scanned-policy.pdf")
    (output / "manifest.json").write_text(
        json.dumps({"schema": 1, "documents": DOCUMENTS}, indent=2),
        encoding="utf-8",
    )
    print(output)


def make_markdown(path: Path) -> None:
    path.write_text(
        """# Structured release policy

## Release gate

The structured policy release code is **MD-4821**. A release requires an
immutable evaluation dataset, an artifact hash and a successful rollback drill.

## Evidence

Answers must cite the document that directly contains the requested fact.
""",
        encoding="utf-8",
    )


def make_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Operations recovery matrix", level=1)
    document.add_paragraph(
        "This Word document verifies table extraction and semantic chunking."
    )
    table = document.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Exercise"
    table.rows[0].cells[1].text = "Frequency"
    table.rows[0].cells[2].text = "Verification code"
    values = [
        ("Object storage restore", "Quarterly", "MINIO-1048"),
        ("Database recovery", "Monthly", "DOCX-5932"),
        ("Vector index rebuild", "Before model migration", "VECTOR-2207"),
    ]
    for exercise, frequency, code in values:
        cells = table.add_row().cells
        cells[0].text = exercise
        cells[1].text = frequency
        cells[2].text = code
    document.save(path)


def make_native_pdf(path: Path) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.setTitle("Native policy fixture")
    text = pdf.beginText(72, 760)
    text.setFont("Helvetica-Bold", 16)
    text.textLine("Native PDF audit policy")
    text.setFont("Helvetica", 12)
    text.textLine("")
    text.textLine("The native PDF audit code is PDF-6843.")
    text.textLine("This page must retain page number 1 in citation metadata.")
    pdf.drawText(text)
    pdf.showPage()
    pdf.save()


def make_scanned_pdf(path: Path) -> None:
    image = Image.new("RGB", (1654, 2339), "white")
    draw = ImageDraw.Draw(image)
    try:
        heading = ImageFont.truetype("arial.ttf", 58)
        body = ImageFont.truetype("arial.ttf", 42)
    except OSError:
        heading = ImageFont.load_default()
        body = ImageFont.load_default()
    draw.text((140, 220), "SCANNED EMERGENCY POLICY", fill="black", font=heading)
    draw.text(
        (140, 390),
        "The scanned document emergency code is OCR-7319.",
        fill="black",
        font=body,
    )
    draw.text(
        (140, 480),
        "This PDF contains only pixels and requires OCR.",
        fill="black",
        font=body,
    )
    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.drawImage(
        ImageReader(image), 0, 0, width=A4[0], height=A4[1], preserveAspectRatio=True
    )
    pdf.showPage()
    pdf.save()


if __name__ == "__main__":
    main()
